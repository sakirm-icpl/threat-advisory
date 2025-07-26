from flask import Blueprint, request, jsonify, send_file, make_response, Response
from app.models.vendor import Vendor
from app.models.product import Product
from app.models.detection_method import DetectionMethod
from app.models.setup_guide import SetupGuide
from app.services.auth import require_admin
from app import db
import json
import csv
import io
from datetime import datetime, timezone
import tempfile
import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os
from sqlalchemy import func, and_, or_
from collections import defaultdict
from collections import OrderedDict
import sys
import traceback
from copy import deepcopy

# Helper: Deeply filter an object to only allowed keys (structure-preserving)
def filter_json_by_sample(input_obj, sample_obj):
    if isinstance(sample_obj, list) and isinstance(input_obj, list):
        return [filter_json_by_sample(item, sample_obj[0]) for item in input_obj]
    elif isinstance(sample_obj, dict) and isinstance(input_obj, dict):
        filtered = {}
        for key in sample_obj:
            if key in input_obj:
                filtered[key] = filter_json_by_sample(input_obj[key], sample_obj[key])
        return filtered
    else:
        return input_obj

# Sample JSON structure for validation and filtering
SAMPLE_JSON = {
    'vendor': {
        'name': '',
        'products': [
            {
                'name': '',
                'category': '',
                'description': '',
                'detection_methods': [
                    {
                        'name': '',
                        'technique': '',
                        'regex_python': '',
                        'regex_ruby': '',
                        'curl_command': '',
                        'expected_response': '',
                        'requires_auth': False
                    }
                ],
                'setup_guides': [
                    {
                        'title': '',
                        'content': ''
                    }
                ]
            }
        ]
    }
}

# Helper: Deep compare two dicts/lists
import collections.abc

def deep_equal(a, b):
    if isinstance(a, dict) and isinstance(b, dict):
        if set(a.keys()) != set(b.keys()):
            return False
        return all(deep_equal(a[k], b[k]) for k in a)
    elif isinstance(a, list) and isinstance(b, list):
        if len(a) != len(b):
            return False
        # Compare lists as sets of dicts (order-insensitive)
        a_sorted = sorted(a, key=lambda x: str(x))
        b_sorted = sorted(b, key=lambda x: str(x))
        return all(deep_equal(x, y) for x, y in zip(a_sorted, b_sorted))
    else:
        return a == b

bp = Blueprint('bulk', __name__, url_prefix='/bulk')





@bp.route('/cleanup', methods=['POST'])
@require_admin
def cleanup_data():
    """Clean up orphaned records and fix data integrity issues"""
    try:
        data = request.json or {}
        cleanup_types = data.get('types', ['orphaned_products', 'orphaned_methods', 'orphaned_guides', 'empty_vendors'])
        
        results = {}
        
        if 'orphaned_products' in cleanup_types:
            orphaned_count = Product.query.filter(~Product.vendor_id.in_(
                db.session.query(Vendor.id)
            )).delete()
            results['orphaned_products_removed'] = orphaned_count
        
        if 'orphaned_methods' in cleanup_types:
            orphaned_count = DetectionMethod.query.filter(~DetectionMethod.product_id.in_(
                db.session.query(Product.id)
            )).delete()
            results['orphaned_methods_removed'] = orphaned_count
        
        if 'orphaned_guides' in cleanup_types:
            orphaned_count = SetupGuide.query.filter(~SetupGuide.product_id.in_(
                db.session.query(Product.id)
            )).delete()
            results['orphaned_guides_removed'] = orphaned_count
        
        if 'empty_vendors' in cleanup_types:
            empty_count = Vendor.query.filter(~Vendor.id.in_(
                db.session.query(Product.vendor_id)
            )).delete()
            results['empty_vendors_removed'] = empty_count
        
        db.session.commit()
        
        return jsonify({
            'message': 'Cleanup completed',
            'results': results
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@bp.route('/export-selective', methods=['POST'])
@require_admin
def export_selective():
    """Export specific data types with filters"""
    try:
        data = request.json or {}
        export_types = data.get('types', ['vendors', 'products', 'methods', 'guides'])
        format_type = data.get('format', 'json')
        filters = data.get('filters', {})
        
        export_data = {}
        
        if 'vendors' in export_types:
            query = Vendor.query
            if 'vendor_ids' in filters:
                query = query.filter(Vendor.id.in_(filters['vendor_ids']))
            if 'vendor_names' in filters:
                query = query.filter(Vendor.name.in_(filters['vendor_names']))
            export_data['vendors'] = [vendor.to_dict() for vendor in query.all()]
        
        if 'products' in export_types:
            query = Product.query
            if 'vendor_ids' in filters:
                query = query.filter(Product.vendor_id.in_(filters['vendor_ids']))
            if 'product_ids' in filters:
                query = query.filter(Product.id.in_(filters['product_ids']))
            if 'categories' in filters:
                query = query.filter(Product.category.in_(filters['categories']))
            export_data['products'] = [product.to_dict() for product in query.all()]
        
        if 'methods' in export_types:
            query = DetectionMethod.query
            if 'product_ids' in filters:
                query = query.filter(DetectionMethod.product_id.in_(filters['product_ids']))
            if 'method_ids' in filters:
                query = query.filter(DetectionMethod.id.in_(filters['method_ids']))
            export_data['methods'] = [method.to_dict() for method in query.all()]
        
        if 'guides' in export_types:
            query = SetupGuide.query
            if 'product_ids' in filters:
                query = query.filter(SetupGuide.product_id.in_(filters['product_ids']))
            if 'guide_ids' in filters:
                query = query.filter(SetupGuide.id.in_(filters['guide_ids']))
            export_data['guides'] = [guide.to_dict() for guide in query.all()]
        
        export_data['exported_at'] = datetime.utcnow().isoformat()
        export_data['filters_applied'] = filters
        
        if format_type == 'json':
            return jsonify(export_data), 200
        elif format_type == 'csv':
            # Create CSV with selected data
            output = io.StringIO()
            writer = csv.writer(output)
            
            if 'vendors' in export_data:
                writer.writerow(['Type', 'ID', 'Name', 'Created At'])
                for vendor in export_data['vendors']:
                    writer.writerow(['Vendor', vendor['id'], vendor['name'], vendor['created_at']])
            
            if 'products' in export_data:
                for product in export_data['products']:
                    writer.writerow(['Product', product['id'], product['name'], product['category'], product['vendor_name']])
            
            output.seek(0)
            return send_file(
                io.BytesIO(output.getvalue().encode('utf-8')),
                mimetype='text/csv',
                as_attachment=True,
                download_name=f'versionintel_selective_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
            )
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def flatten_import_json(nested_json):
    """Flatten nested vendor->products->methods/guides JSON into flat lists for import."""
    vendors = []
    products = []
    methods = []
    guides = []
    # Support both {vendor: {...}} and {vendors: [...]}
    vendor_objs = []
    if 'vendor' in nested_json:
        vendor_objs.append(nested_json['vendor'])
    if 'vendors' in nested_json:
        vendor_objs.extend(nested_json['vendors'])
    for v in vendor_objs:
        v_name = v.get('name', '').strip()
        if not v_name:
            continue
        vendors.append({'name': v_name})
        for p in v.get('products', []):
            p_name = p.get('name', '').strip()
            if not p_name:
                continue
            products.append({
                'name': p_name,
                'vendor': v_name,
                'category': p.get('category', ''),
                'description': p.get('description', '')
            })
            for m in p.get('detection_methods', []):
                methods.append({
                    'name': m.get('name', ''),
                    'product': p_name,
                    'technique': m.get('technique', ''),
                    'regex_python': m.get('regex_python', ''),
                    'regex_ruby': m.get('regex_ruby', ''),
                    'curl_command': m.get('curl_command', ''),
                    'expected_response': m.get('expected_response', ''),
                    'requires_auth': m.get('requires_auth', False)
                })
            for g in p.get('setup_guides', []):
                guides.append({
                    'title': g.get('title', ''),
                    'product': p_name,
                    'content': g.get('content', '')
                })
    return {
        'vendors': vendors,
        'products': products,
        'methods': methods,
        'guides': guides
    }

@bp.route('/import-preview', methods=['POST'])
@require_admin
def import_preview():
    """Preview what will be imported or replaced, with deep comparison and structure cleaning."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        import_data = data.get('data', data)
        # Clean/filter the uploaded JSON to match the sample structure
        cleaned_json = filter_json_by_sample(import_data, SAMPLE_JSON)
        # Flatten for easier DB comparison
        flat = flatten_import_json(cleaned_json)
        # DB snapshot
        vendor_name_to_obj = {v.name: v for v in Vendor.query.all()}
        print("[DEBUG] Vendor names in DB:", list(vendor_name_to_obj.keys()))
        vendor = cleaned_json.get('vendor')
        print("[DEBUG] Incoming vendor name:", vendor.get('name') if vendor else None)
        product_name_to_obj = {p.name: p for p in Product.query.all()}
        method_key_to_obj = {(m.name, m.product_id): m for m in DetectionMethod.query.all()}
        guide_key_to_obj = {(g.title if hasattr(g, 'title') else '', g.product_id): g for g in SetupGuide.query.all()}
        # --- Deep comparison logic ---
        preview = {'can_add': False, 'can_replace': False, 'add': {}, 'replace': {}, 'error': None}
        vendor = cleaned_json.get('vendor')
        if not vendor or not vendor.get('name'):
            preview['error'] = 'Vendor name is required.'
            return jsonify(preview), 400
        v_name = vendor['name']
        db_vendor = vendor_name_to_obj.get(v_name)
        if not db_vendor:
            print("[DEBUG] Vendor does not exist, can add everything.")
            preview['can_add'] = True
            preview['add'] = cleaned_json
            print("[DEBUG] Returning preview:", preview)
            return jsonify(preview), 200
        # Vendor exists, check products
        db_products = [p for p in Product.query.filter_by(vendor_id=db_vendor.id).all()]
        db_product_names = {p.name for p in db_products}
        add_products = []
        replace_products = []
        all_products_unchanged = True  # Track if all products are unchanged
        for prod in vendor.get('products', []):
            p_name = prod.get('name')
            if not p_name:
                continue
            db_prod = next((p for p in db_products if p.name == p_name), None)
            if not db_prod:
                add_products.append(prod)
                all_products_unchanged = False
                continue
            # Product exists, deep compare
            db_methods = [m for m in DetectionMethod.query.filter_by(product_id=db_prod.id).all()]
            db_guides = [g for g in SetupGuide.query.filter_by(product_id=db_prod.id).all()]
            prod_methods = prod.get('detection_methods', [])
            db_methods_dict = {(m.name): m for m in db_methods}
            print(f"[DEBUG] Product '{p_name}' has {len(db_methods)} existing methods: {[m.name for m in db_methods]}")
            print(f"[DEBUG] Import file has {len(prod_methods)} methods: {[m.get('name') for m in prod_methods]}")
            add_methods = []
            replace_methods = []
            for m in prod_methods:
                m_name = m.get('name')
                if not m_name:
                    continue
                print(f"[DEBUG] Checking method '{m_name}'")
                # Try to match by all fields except name
                found_equivalent = False
                for db_m in db_methods:
                    db_m_dict = {k: getattr(db_m, k) for k in ['technique','regex_python','regex_ruby','curl_command','expected_response','requires_auth']}
                    m_compare = {k: m.get(k) for k in ['technique','regex_python','regex_ruby','curl_command','expected_response','requires_auth']}
                    print(f"[DEBUG] Comparing method '{m_name}' with DB method '{db_m.name}':")
                    print(f"[DEBUG]   Import: {m_compare}")
                    print(f"[DEBUG]   DB:     {db_m_dict}")
                    print(f"[DEBUG]   Deep equal: {deep_equal(m_compare, db_m_dict)}")
                    if deep_equal(m_compare, db_m_dict):
                        found_equivalent = True
                        print(f"[DEBUG] Found equivalent method in DB for '{m_name}' (ignoring name)")
                        break
                if not found_equivalent:
                    print(f"[DEBUG] No equivalent method found for '{m_name}', adding to add_methods")
                    add_methods.append(m)
                    all_products_unchanged = False
                else:
                    print(f"[DEBUG] Method '{m_name}' matches an existing method (ignoring name)")
            prod_guides = prod.get('setup_guides', [])
            add_guides = []
            replace_guides = []
            for g in prod_guides:
                g_title = g.get('title')
                g_content = g.get('content', '')
                if not g_title or not g_content:
                    continue
                print(f"[DEBUG] Checking guide '{g_title}'")
                # Try to match by content (ignoring title)
                found_equivalent = False
                for db_g in db_guides:
                    print(f"[DEBUG] Comparing guide '{g_title}' with DB guide:")
                    print(f"[DEBUG]   Import content: '{g_content}'")
                    print(f"[DEBUG]   DB instructions: '{db_g.instructions}'")
                    print(f"[DEBUG]   Direct match: {db_g.instructions == g_content}")
                    print(f"[DEBUG]   Title format match: {db_g.instructions == f'Title: {g_title}\\n\\n{g_content}'}")
                    # Check if content matches (ignoring title format)
                    if db_g.instructions == g_content or db_g.instructions == f"Title: {g_title}\n\n{g_content}":
                        found_equivalent = True
                        print(f"[DEBUG] Found equivalent guide in DB for '{g_title}' (ignoring title)")
                        break
                if not found_equivalent:
                    print(f"[DEBUG] No equivalent guide found for '{g_title}', adding to add_guides")
                    add_guides.append(g)
                    all_products_unchanged = False
                else:
                    print(f"[DEBUG] Guide '{g_title}' matches an existing guide (ignoring title)")
            print(f"[DEBUG] Product '{p_name}' - add_methods: {len(add_methods)}, replace_methods: {len(replace_methods)}, add_guides: {len(add_guides)}, replace_guides: {len(replace_guides)}")
            if not add_methods and not replace_methods and not add_guides and not replace_guides:
                print(f"[DEBUG] Product '{p_name}' has no changes, skipping")
                continue
            print(f"[DEBUG] Product '{p_name}' has changes, adding to replace_products")
            replace_products.append({
                'name': p_name,
                'add_methods': add_methods,
                'replace_methods': replace_methods,
                'add_guides': add_guides,
                'replace_guides': replace_guides
            })
        # If all products exist and match, error
        print("[DEBUG] Final check - add_products:", len(add_products), "all_products_unchanged:", all_products_unchanged)
        if not add_products and all_products_unchanged:
            print("[DEBUG] All products exist and match. Returning error.")
            preview['error'] = 'The data already exists.'
            print("[DEBUG] Returning preview:", preview)
            return jsonify(preview), 200
        # If there are products to add, can add
        if add_products:
            print("[DEBUG] There are products to add.")
            preview['can_add'] = True
            preview['add'] = {'vendor': {'name': v_name, 'products': add_products}}
        # Only set can_replace if there are actual differences to replace
        filtered_replace_products = [p for p in replace_products if (
            (p.get('replace_methods') and len(p['replace_methods']) > 0) or
            (p.get('replace_guides') and len(p['replace_guides']) > 0)
        )]
        if filtered_replace_products:
            print("[DEBUG] There are products to replace.")
            preview['can_replace'] = True
            preview['replace'] = {'vendor': {'name': v_name, 'products': filtered_replace_products}}
        print("[DEBUG] Returning preview:", preview)
        return jsonify(preview), 200
    except Exception as e:
        import sys, traceback
        print("IMPORT PREVIEW ERROR:", file=sys.stderr, flush=True)
        print(traceback.format_exc(), file=sys.stderr, flush=True)
        return jsonify({'error': str(e)}), 500

@bp.route('/bulk-delete', methods=['POST'])
@require_admin
def bulk_delete():
    """Safely delete multiple records with confirmation"""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        delete_types = data.get('types', {})
        results = {}
        
        if 'vendor_ids' in delete_types:
            vendor_ids = delete_types['vendor_ids']
            # Delete vendors and all related data
            vendors = Vendor.query.filter(Vendor.id.in_(vendor_ids)).all()
            deleted_count = 0
            for vendor in vendors:
                # Delete related products, methods, and guides
                products = Product.query.filter_by(vendor_id=vendor.id).all()
                for product in products:
                    DetectionMethod.query.filter_by(product_id=product.id).delete()
                    SetupGuide.query.filter_by(product_id=product.id).delete()
                Product.query.filter_by(vendor_id=vendor.id).delete()
                db.session.delete(vendor)
                deleted_count += 1
            results['vendors_deleted'] = deleted_count
        
        if 'product_ids' in delete_types:
            product_ids = delete_types['product_ids']
            # Delete products and related methods/guides
            products = Product.query.filter(Product.id.in_(product_ids)).all()
            deleted_count = 0
            for product in products:
                DetectionMethod.query.filter_by(product_id=product.id).delete()
                SetupGuide.query.filter_by(product_id=product.id).delete()
                db.session.delete(product)
                deleted_count += 1
            results['products_deleted'] = deleted_count
        
        if 'method_ids' in delete_types:
            method_ids = delete_types['method_ids']
            deleted_count = DetectionMethod.query.filter(DetectionMethod.id.in_(method_ids)).delete()
            results['methods_deleted'] = deleted_count
        
        if 'guide_ids' in delete_types:
            guide_ids = delete_types['guide_ids']
            deleted_count = SetupGuide.query.filter(SetupGuide.id.in_(guide_ids)).delete()
            results['guides_deleted'] = deleted_count
        
        db.session.commit()
        
        return jsonify({
            'message': 'Bulk delete completed',
            'results': results
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500



@bp.route('/export', methods=['GET'])
@require_admin
def export_data():
    """Export all data as JSON"""
    try:
        format_type = request.args.get('format', 'json')  # json, csv
        
        if format_type == 'json':
            data = {
                'vendors': [vendor.to_dict() for vendor in Vendor.query.all()],
                'products': [product.to_dict() for product in Product.query.all()],
                'methods': [method.to_dict() for method in DetectionMethod.query.all()],
                'guides': [guide.to_dict() for guide in SetupGuide.query.all()],
                'exported_at': datetime.utcnow().isoformat()
            }
            
            return jsonify(data), 200
        
        elif format_type == 'csv':
            # Create CSV files for each entity
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Export vendors
            writer.writerow(['Vendor ID', 'Name', 'Created At', 'Updated At'])
            for vendor in Vendor.query.all():
                writer.writerow([vendor.id, vendor.name, vendor.created_at, vendor.updated_at])
            
            output.seek(0)
            return send_file(
                io.BytesIO(output.getvalue().encode('utf-8')),
                mimetype='text/csv',
                as_attachment=True,
                download_name=f'versionintel_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
            )
        
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/import', methods=['POST'])
@require_admin
def import_data():
    """Import data from JSON with proper handling of nested structure and add/replace modes."""
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        import_data = data.get('data', data)
        mode = data.get('mode', 'add')  # 'add' or 'replace'
        
        # Clean/filter the uploaded JSON to match the sample structure
        cleaned_json = filter_json_by_sample(import_data, SAMPLE_JSON)
        
        summary = {
            'vendors': {'added': [], 'errors': []},
            'products': {'added': [], 'errors': []},
            'methods': {'added': [], 'errors': []},
            'guides': {'added': [], 'errors': []},
            'errors': []
        }

        vendor_data = cleaned_json.get('vendor')
        if not vendor_data or not vendor_data.get('name'):
            return jsonify({'error': 'Vendor name is required'}), 400

        vendor_name = vendor_data['name'].strip()
        
        # Check if vendor exists
        existing_vendor = Vendor.query.filter_by(name=vendor_name).first()
        
        if not existing_vendor:
            # Create new vendor
            try:
                existing_vendor = Vendor(name=vendor_name)
                db.session.add(existing_vendor)
                db.session.flush()  # Get the ID
                summary['vendors']['added'].append({'name': vendor_name})
            except Exception as e:
                summary['vendors']['errors'].append(f"Vendor {vendor_name}: {str(e)}")
                db.session.rollback()
                return jsonify({'error': f"Failed to create vendor: {str(e)}"}), 500

        # Process products
        for product_data in vendor_data.get('products', []):
            product_name = product_data.get('name', '').strip()
            if not product_name:
                summary['products']['errors'].append('Product with missing name')
                continue

            # Check if product exists
            existing_product = Product.query.filter_by(name=product_name, vendor_id=existing_vendor.id).first()
            
            if not existing_product:
                # Create new product
                try:
                    existing_product = Product(
                        name=product_name,
                        vendor_id=existing_vendor.id,
                        category=product_data.get('category'),
                        description=product_data.get('description')
                    )
                    db.session.add(existing_product)
                    db.session.flush()  # Get the ID
                    summary['products']['added'].append({'name': product_name, 'vendor': vendor_name})
                except Exception as e:
                    summary['products']['errors'].append(f"Product {product_name}: {str(e)}")
                    continue

            # Process detection methods
            for method_data in product_data.get('detection_methods', []):
                method_name = method_data.get('name', '').strip()
                if not method_name:
                    summary['methods']['errors'].append(f"Method missing name for product {product_name}")
                    continue

                existing_method = DetectionMethod.query.filter_by(name=method_name, product_id=existing_product.id).first()
                
                if not existing_method:
                    # Add new method
                    try:
                        method = DetectionMethod(
                            name=method_name,
                            product_id=existing_product.id,
                            technique=method_data.get('technique'),
                            regex_python=method_data.get('regex_python'),
                            regex_ruby=method_data.get('regex_ruby'),
                            curl_command=method_data.get('curl_command'),
                            expected_response=method_data.get('expected_response'),
                            requires_auth=method_data.get('requires_auth', False)
                        )
                        db.session.add(method)
                        summary['methods']['added'].append({'name': method_name, 'product': product_name})
                    except Exception as e:
                        summary['methods']['errors'].append(f"Method {method_name}: {str(e)}")
                elif mode == 'replace':
                    # Replace existing method
                    try:
                        existing_method.technique = method_data.get('technique')
                        existing_method.regex_python = method_data.get('regex_python')
                        existing_method.regex_ruby = method_data.get('regex_ruby')
                        existing_method.curl_command = method_data.get('curl_command')
                        existing_method.expected_response = method_data.get('expected_response')
                        existing_method.requires_auth = method_data.get('requires_auth', False)
                        summary['methods']['added'].append({'name': method_name, 'product': product_name, 'action': 'replaced'})
                    except Exception as e:
                        summary['methods']['errors'].append(f"Method {method_name}: {str(e)}")

            # Process setup guides
            for guide_data in product_data.get('setup_guides', []):
                guide_title = guide_data.get('title', '').strip()
                if not guide_title:
                    summary['guides']['errors'].append(f"Guide missing title for product {product_name}")
                    continue

                # For SetupGuide, we'll use the title as part of the instructions since there's no title field
                guide_content = guide_data.get('content', '')
                if not guide_content:
                    summary['guides']['errors'].append(f"Guide missing content for product {product_name}")
                    continue

                # Check if guide exists by comparing content (since no title field)
                existing_guide = SetupGuide.query.filter_by(
                    product_id=existing_product.id,
                    instructions=guide_content
                ).first()
                
                if not existing_guide:
                    # Add new guide
                    try:
                        guide = SetupGuide(
                            instructions=f"Title: {guide_title}\n\n{guide_content}",
                            product_id=existing_product.id
                        )
                        db.session.add(guide)
                        summary['guides']['added'].append({'title': guide_title, 'product': product_name})
                    except Exception as e:
                        summary['guides']['errors'].append(f"Guide {guide_title}: {str(e)}")
                elif mode == 'replace':
                    # Replace existing guide (update content)
                    try:
                        existing_guide.instructions = f"Title: {guide_title}\n\n{guide_content}"
                        summary['guides']['added'].append({'title': guide_title, 'product': product_name, 'action': 'replaced'})
                    except Exception as e:
                        summary['guides']['errors'].append(f"Guide {guide_title}: {str(e)}")

        db.session.commit()
        return jsonify(summary), 200
        
    except Exception as e:
        db.session.rollback()
        import sys, traceback
        print("IMPORT ERROR (POST /bulk/import):", file=sys.stderr, flush=True)
        print(traceback.format_exc(), file=sys.stderr, flush=True)
        return jsonify({'error': str(e)}), 500

@bp.route('/backup', methods=['GET'])
@require_admin
def create_backup():
    """Create a complete backup of all data"""
    try:
        backup_data = {
            'metadata': {
                'created_at': datetime.utcnow().isoformat(),
                'version': '1.0.0',
                'description': 'VersionIntel complete backup'
            },
            'data': {
                'vendors': [vendor.to_dict() for vendor in Vendor.query.all()],
                'products': [product.to_dict() for product in Product.query.all()],
                'methods': [method.to_dict() for method in DetectionMethod.query.all()],
                'guides': [guide.to_dict() for guide in SetupGuide.query.all()]
            }
        }
        
        return jsonify(backup_data), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@bp.route('/restore', methods=['POST'])
@require_admin
def restore_backup():
    """Restore data from backup"""
    try:
        backup_data = request.json
        if not backup_data or 'data' not in backup_data:
            return jsonify({'error': 'Invalid backup data'}), 400
        
        # Clear existing data
        DetectionMethod.query.delete()
        SetupGuide.query.delete()
        Product.query.delete()
        Vendor.query.delete()
        
        data = backup_data['data']
        restored = {'vendors': 0, 'products': 0, 'methods': 0, 'guides': 0}
        
        # Restore vendors
        for vendor_data in data.get('vendors', []):
            vendor = Vendor(name=vendor_data['name'])
            db.session.add(vendor)
            restored['vendors'] += 1
        
        # Restore products
        for product_data in data.get('products', []):
            product = Product(
                name=product_data['name'],
                vendor_id=product_data['vendor_id'],
                category=product_data.get('category'),
                description=product_data.get('description')
            )
            db.session.add(product)
            restored['products'] += 1
        
        # Restore methods
        for method_data in data.get('methods', []):
            method = DetectionMethod(**method_data)
            db.session.add(method)
            restored['methods'] += 1
        
        # Restore guides
        for guide_data in data.get('guides', []):
            guide = SetupGuide(**guide_data)
            db.session.add(guide)
            restored['guides'] += 1
        
        db.session.commit()
        
        return jsonify({
            'message': 'Restore completed',
            'restored': restored
        }), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500 

def isoformat_z(dt):
    if not dt:
        return ""
    if isinstance(dt, str):
        return dt
    try:
        return dt.astimezone(timezone.utc).isoformat().replace('+00:00', 'Z')
    except Exception:
        return str(dt)

def clean_method(method):
    return OrderedDict([
        ('name', str(method.name) if method.name else ""),
        ('technique', str(method.technique) if method.technique else ""),
        ('regex_python', str(method.regex_python) if method.regex_python else ""),
        ('regex_ruby', str(method.regex_ruby) if method.regex_ruby else ""),
        ('curl_command', str(method.curl_command) if method.curl_command else ""),
        ('expected_response', str(method.expected_response) if method.expected_response else ""),
        ('requires_auth', bool(method.requires_auth)),
        ('created_at', isoformat_z(getattr(method, 'created_at', None))),
        ('updated_at', isoformat_z(getattr(method, 'updated_at', None)))
    ])

def clean_guide(guide, product_name=None):
    title = getattr(guide, 'title', None)
    if not title or str(title).strip() == "" or title == 'null':
        title = product_name or "Untitled"
    content = getattr(guide, 'instructions', None) or getattr(guide, 'content', None) or ""
    return OrderedDict([
        ('title', str(title)),
        ('content', str(content)),
        ('created_at', isoformat_z(getattr(guide, 'created_at', None))),
        ('updated_at', isoformat_z(getattr(guide, 'updated_at', None)))
    ])

def clean_product(product):
    methods = DetectionMethod.query.filter_by(product_id=product.id).all()
    guides = SetupGuide.query.filter_by(product_id=product.id).all()
    return OrderedDict([
        ('name', str(product.name) if product.name else ""),
        ('category', str(product.category) if product.category else ""),
        ('description', str(product.description) if product.description else ""),
        ('created_at', isoformat_z(getattr(product, 'created_at', None))),
        ('updated_at', isoformat_z(getattr(product, 'updated_at', None))),
        ('detection_methods', [clean_method(m) for m in methods]),
        ('setup_guides', [clean_guide(g, product.name) for g in guides])
    ])

def clean_vendor(vendor):
    products = Product.query.filter_by(vendor_id=vendor.id).all()
    return OrderedDict([
        ('name', str(vendor.name) if vendor.name else ""),
        ('created_at', isoformat_z(getattr(vendor, 'created_at', None))),
        ('updated_at', isoformat_z(getattr(vendor, 'updated_at', None))),
        ('products', [clean_product(p) for p in products])
    ])

def safe_json_dumps(obj):
    def fallback(o):
        return str(o)
    return json.dumps(obj, ensure_ascii=False, default=fallback)

@bp.route('/export-all', methods=['GET'])
@require_admin
def export_all():
    vendors = Vendor.query.all()
    export_data = OrderedDict([
        ('exported_at', isoformat_z(datetime.utcnow())),
        ('vendors', [clean_vendor(v) for v in vendors])
    ])
    json_str = json.dumps(export_data, ensure_ascii=False)
    return Response(json_str, mimetype='application/json')

@bp.route('/export-vendor/<int:vendor_id>', methods=['GET'])
@require_admin
def export_vendor_data(vendor_id):
    vendor = Vendor.query.get(vendor_id)
    if not vendor:
        return jsonify({'error': 'Vendor not found'}), 404
    export_data = OrderedDict([
        ('exported_at', isoformat_z(datetime.utcnow())),
        ('vendor', clean_vendor(vendor))
    ])
    json_str = json.dumps(export_data, ensure_ascii=False)
    return Response(json_str, mimetype='application/json')

@bp.route('/export-all-complete', methods=['GET'])
@require_admin
def export_all_complete():
    """Export all data in a complete format with full relationships"""
    try:
        format_type = request.args.get('format', 'json')
        
        # Get all vendors with their complete data
        vendors = Vendor.query.all()
        complete_data = []
        
        for vendor in vendors:
            vendor_dict = vendor.to_dict()
            products = Product.query.filter_by(vendor_id=vendor.id).all()
            vendor_dict['products'] = []
            
            for product in products:
                product_dict = product.to_dict()
                
                # Add detection methods for this product
                methods = DetectionMethod.query.filter_by(product_id=product.id).all()
                product_dict['detection_methods'] = [m.to_dict() for m in methods]
                
                # Add setup guides for this product
                guides = SetupGuide.query.filter_by(product_id=product.id).all()
                product_dict['setup_guides'] = [g.to_dict() for g in guides]
                
                vendor_dict['products'].append(product_dict)
            
            complete_data.append(vendor_dict)
        
        export_data = {
            'vendors': complete_data,
            'exported_at': datetime.utcnow().isoformat(),
            'summary': {
                'total_vendors': len(vendors),
                'total_products': sum(len(v['products']) for v in complete_data),
                'total_methods': sum(len(p['detection_methods']) for v in complete_data for p in v['products']),
                'total_guides': sum(len(p['setup_guides']) for v in complete_data for p in v['products'])
            }
        }
        
        if format_type == 'json':
            return jsonify(export_data), 200
        elif format_type == 'csv':
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Write summary
            writer.writerow(['Export Summary'])
            writer.writerow(['Total Vendors', export_data['summary']['total_vendors']])
            writer.writerow(['Total Products', export_data['summary']['total_products']])
            writer.writerow(['Total Methods', export_data['summary']['total_methods']])
            writer.writerow(['Total Guides', export_data['summary']['total_guides']])
            writer.writerow([])
            
            # Write detailed data
            writer.writerow(['Vendor', 'Product', 'Category', 'Description', 'Detection Methods', 'Setup Guides'])
            
            for vendor in complete_data:
                for product in vendor['products']:
                    method_names = ', '.join([m.get('name', '') for m in product['detection_methods']])
                    guide_titles = ', '.join([g.get('title', '') for g in product['setup_guides']])
                    
                    writer.writerow([
                        vendor['name'],
                        product['name'],
                        product.get('category', ''),
                        product.get('description', ''),
                        method_names,
                        guide_titles
                    ])
            
            output.seek(0)
            return send_file(
                io.BytesIO(output.getvalue().encode('utf-8')),
                mimetype='text/csv',
                as_attachment=True,
                download_name=f'versionintel_complete_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
            )
        elif format_type == 'docx':
            try:
                # Create DOCX with complete data
                doc = Document()
                
                # Add title
                title = doc.add_heading('VersionIntel Complete Export', 0)
                title.alignment = 1  # Center alignment
                
                # Add export timestamp
                timestamp_para = doc.add_paragraph(f'Exported on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                timestamp_para.alignment = 1  # Center alignment
                doc.add_paragraph()
                
                # Add summary section
                doc.add_heading('Export Summary', level=1)
                summary_table = doc.add_table(rows=1, cols=2)
                summary_table.style = 'Table Grid'
                hdr_cells = summary_table.rows[0].cells
                hdr_cells[0].text = 'Metric'
                hdr_cells[1].text = 'Count'
                
                # Style the header row
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].bold = True
                
                for key, value in export_data['summary'].items():
                    row_cells = summary_table.add_row().cells
                    row_cells[0].text = key.replace('_', ' ').title()
                    row_cells[1].text = str(value)
                
                doc.add_paragraph()
                
                # Add vendor data
                for vendor in complete_data:
                    vendor_heading = doc.add_heading(f'Vendor: {vendor["name"]}', level=1)
                    
                    if vendor['products']:
                        for product in vendor['products']:
                            product_heading = doc.add_heading(f'Product: {product["name"]}', level=2)
                            
                            # Product details in a table
                            product_table = doc.add_table(rows=0, cols=2)
                            product_table.style = 'Table Grid'
                            
                            # Add category
                            if product.get('category'):
                                row = product_table.add_row().cells
                                row[0].text = 'Category'
                                row[1].text = product['category']
                            
                            # Add description
                            if product.get('description'):
                                row = product_table.add_row().cells
                                row[0].text = 'Description'
                                row[1].text = product['description']
                            
                            doc.add_paragraph()
                            
                            # Detection methods
                            if product['detection_methods']:
                                doc.add_heading('Detection Methods', level=3)
                                method_table = doc.add_table(rows=1, cols=2)
                                method_table.style = 'Table Grid'
                                method_hdr = method_table.rows[0].cells
                                method_hdr[0].text = 'Method Name'
                                method_hdr[1].text = 'Technique'
                                
                                for method in product['detection_methods']:
                                    method_row = method_table.add_row().cells
                                    method_row[0].text = method.get('name', 'N/A')
                                    method_row[1].text = method.get('technique', 'N/A')
                                
                                doc.add_paragraph()
                            
                            # Setup guides
                            if product['setup_guides']:
                                doc.add_heading('Setup Guides', level=3)
                                guide_table = doc.add_table(rows=1, cols=2)
                                guide_table.style = 'Table Grid'
                                guide_hdr = guide_table.rows[0].cells
                                guide_hdr[0].text = 'Title'
                                guide_hdr[1].text = 'Content'
                                
                                for guide in product['setup_guides']:
                                    guide_row = guide_table.add_row().cells
                                    guide_row[0].text = guide.get('title', 'Untitled')
                                    guide_row[1].text = guide.get('content', 'No content available')
                                
                                doc.add_paragraph()
                            
                            doc.add_paragraph()
                    else:
                        doc.add_paragraph('No products found for this vendor.')
                        doc.add_paragraph()
                
                # Save to BytesIO instead of temporary file
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                
                return send_file(
                    docx_buffer,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f'versionintel_complete_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                )
                
            except Exception as docx_error:
                return jsonify({'error': f'DOCX generation failed: {str(docx_error)}'}), 500
            
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
    except Exception as e:
        return jsonify({'error': str(e)}), 500 

@bp.route('/export-product/<int:product_id>', methods=['GET'])
@require_admin
def export_product_data(product_id):
    product = Product.query.get(product_id)
    if not product:
        return jsonify({'error': 'Product not found'}), 404
    vendor = Vendor.query.get(product.vendor_id)
    if not vendor:
        return jsonify({'error': 'Vendor not found'}), 404
    vendor_data = OrderedDict([
        ('name', vendor.name),
        ('created_at', isoformat_z(getattr(vendor, 'created_at', None))),
        ('updated_at', isoformat_z(getattr(vendor, 'updated_at', None))),
        ('products', [clean_product(product)])
    ])
    export_data = OrderedDict([
        ('exported_at', isoformat_z(datetime.utcnow())),
        ('vendor', vendor_data)
    ])
    json_str = json.dumps(export_data, ensure_ascii=False)
    return Response(json_str, mimetype='application/json') 